from __future__ import annotations

import os
import re
import threading
from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm, inch, mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    HRFlowable,
    NextPageTemplate,
    PageBreak,
    PageTemplate,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from reportlab.platypus.flowables import Flowable

from ..models import BookProject

# ── Font registration ──────────────────────────────────────────────────────────

_FONTS_REGISTERED: dict[str, bool] = {}
_FONT_REG_LOCK = threading.Lock()

_STATIC_DIR = os.path.normpath(os.path.join(os.path.dirname(__file__), "..", "static", "fonts"))


def _font_search_dirs() -> list[str]:
    dirs = [
        _STATIC_DIR,
        r"C:\Windows\Fonts",
        "/usr/share/fonts/truetype/dejavu",
        "/usr/share/fonts/dejavu",
        "/usr/share/fonts/truetype/liberation",
        "/usr/share/fonts/liberation",
        "/usr/share/fonts/truetype/freefont",
        "/usr/share/fonts/truetype/noto",
        "/usr/share/fonts/truetype",
        "/usr/share/fonts",
        "/Library/Fonts",
        "/System/Library/Fonts",
        os.path.expanduser("~/Library/Fonts"),
        "/usr/local/share/fonts",
        "/opt/homebrew/share/fonts",
    ]
    return [d for d in dirs if d]


_FONT_SETS = [
    ("Georgia",     "Georgia.ttf",                  "Georgia Bold.ttf",            "Georgia Italic.ttf"),
    ("DejaVu",      "DejaVuSerif.ttf",             "DejaVuSerif-Bold.ttf",        "DejaVuSerif-Italic.ttf"),
    ("DejaVuSans",  "DejaVuSans.ttf",              "DejaVuSans-Bold.ttf",         "DejaVuSans-Oblique.ttf"),
    ("Liberation",  "LiberationSerif-Regular.ttf", "LiberationSerif-Bold.ttf",    "LiberationSerif-Italic.ttf"),
    ("LiberSans",   "LiberationSans-Regular.ttf",  "LiberationSans-Bold.ttf",     "LiberationSans-Italic.ttf"),
    ("FreeSans",    "FreeSans.ttf",                "FreeSansBold.ttf",            "FreeSansOblique.ttf"),
    ("NotoSans",    "NotoSans-Regular.ttf",        "NotoSans-Bold.ttf",           "NotoSans-Italic.ttf"),
]

_FONT_ALIASES: dict[str, str] = {alias: alias for alias, *_ in _FONT_SETS}


def _register_font_variant(name: str, primary_path: str, fallback_path: str) -> bool:
    for path in (primary_path, fallback_path):
        if not path or not os.path.isfile(path):
            continue
        try:
            pdfmetrics.registerFont(TTFont(name, path))
            return True
        except Exception:
            continue
    return False


def _try_register_set(folder: str, fname: str, regular: str, bold: str, italic: str) -> bool:
    rp = os.path.join(folder, regular)
    if not os.path.isfile(rp):
        return False

    bold_name = fname + "Bold"
    italic_name = fname + "Italic"
    bp = os.path.join(folder, bold)
    ip = os.path.join(folder, italic)

    if not _register_font_variant(fname, rp, rp):
        return False

    # ReportLab may request Bold/Italic aliases when rendering Paragraph markup
    # or when the page decorator switches variants. Ensure those aliases always
    # exist, even if the real bold/italic TTF is missing in the container.
    _register_font_variant(bold_name, bp, rp)
    _register_font_variant(italic_name, ip, rp)
    pdfmetrics.registerFontFamily(
        fname,
        normal=fname,
        bold=bold_name,
        italic=italic_name,
        boldItalic=bold_name,
    )
    return True


def _register_fonts():
    """Register all available font families; always marks done."""
    if _FONTS_REGISTERED.get("__done__"):
        return

    with _FONT_REG_LOCK:
        if _FONTS_REGISTERED.get("__done__"):
            return

        for folder in _font_search_dirs():
            if not os.path.isdir(folder):
                continue
            for fname, regular, bold, italic in _FONT_SETS:
                if _FONTS_REGISTERED.get(fname):
                    continue
                if _try_register_set(folder, fname, regular, bold, italic):
                    _FONTS_REGISTERED[fname] = True

        _FONTS_REGISTERED["__done__"] = True


def _set_active_font(family: str) -> tuple[str, str, str]:
    """Return (regular, bold, italic) names for the given family alias."""
    if family == "auto" or family not in _FONTS_REGISTERED:
        # Use first registered family
        for fname, *_ in _FONT_SETS:
            if _FONTS_REGISTERED.get(fname):
                return fname, fname + "Bold", fname + "Italic"
        return "Helvetica", "Helvetica-Bold", "Helvetica-Oblique"
    bold = family + "Bold"
    italic = family + "Italic"
    # Check if bold/italic variants exist
    try:
        pdfmetrics.getFont(bold)
    except Exception:
        bold = family
    try:
        pdfmetrics.getFont(italic)
    except Exception:
        italic = family
    return family, bold, italic


def _coerce_font_family(value: object, default: str = "Georgia") -> str:
    family = str(value or default).strip()
    if family == "auto" or family in _FONT_ALIASES:
        return family
    return default


def _coerce_font_size(value: object, *, default: int, minimum: int, maximum: int) -> int:
    try:
        size = int(value)
    except (TypeError, ValueError):
        return default
    return max(minimum, min(maximum, size))


def _coerce_trim_size(value: object, default: str = "6x9") -> str:
    trim = str(value or default).strip().lower()
    return "6x9" if trim == "6x9" else default


def _trim_page_size(trim_size: str) -> tuple[float, float]:
    if trim_size == "6x9":
        return 6 * inch, 9 * inch
    return A4


def _coerce_bool(value: object, *, default: bool) -> bool:
    if value is None:
        return default
    if isinstance(value, bool):
        return value
    raw = str(value).strip().lower()
    if not raw:
        return default
    return raw in {"1", "true", "yes", "on"}


def _canvas_safe_text(text: object, max_len: int) -> str:
    value = str(text or "")
    # Canvas header/footer text should stay single-line and printable.
    value = re.sub(r"[\x00-\x1F\x7F]+", " ", value)
    value = re.sub(r"\s+", " ", value).strip()
    return value[:max_len]


# ── Color palette ──────────────────────────────────────────────────────────────

INK         = colors.HexColor("#1f1a17")
INK_LIGHT   = colors.HexColor("#52463e")
ACCENT      = colors.HexColor("#8c6a43")
ACCENT_SOFT = colors.HexColor("#d7c4aa")
GOLD        = colors.HexColor("#b58b4a")
CREAM       = colors.HexColor("#f8f3ec")
RULE_COLOR  = colors.HexColor("#d8cec0")
CAPTION_COL = colors.HexColor("#6f655d")
WHITE       = colors.white


# ── Custom flowables ───────────────────────────────────────────────────────────

class ColorRect(Flowable):
    def __init__(self, width, height, fill_color, radius=0):
        super().__init__()
        self.width = width
        self.height = height
        self.fill_color = fill_color
        self.radius = radius

    def draw(self):
        self.canv.setFillColor(self.fill_color)
        if self.radius:
            self.canv.roundRect(0, 0, self.width, self.height, self.radius, fill=1, stroke=0)
        else:
            self.canv.rect(0, 0, self.width, self.height, fill=1, stroke=0)


class OrnamentalRule(Flowable):
    """Decorative chapter rule with central diamond."""
    def __init__(self, width, color=ACCENT):
        super().__init__()
        self.width = width
        self.color = color
        self.height = 12

    def draw(self):
        c = self.canv
        c.setStrokeColor(self.color)
        c.setFillColor(self.color)
        c.setLineWidth(0.8)
        c.line(0, 6, self.width * 0.42, 6)
        cx, cy = self.width / 2, 6
        d = 4
        p = c.beginPath()
        p.moveTo(cx, cy + d)
        p.lineTo(cx + d, cy)
        p.lineTo(cx, cy - d)
        p.lineTo(cx - d, cy)
        p.close()
        c.drawPath(p, fill=1, stroke=0)
        c.line(self.width * 0.58, 6, self.width, 6)


# ── Style builder ──────────────────────────────────────────────────────────────

def _build_styles(
    body_font: str,
    bold_font: str,
    italic_font: str,
    *,
    book_title_size: int = 30,
    chapter_title_size: int = 23,
    subsection_title_size: int = 17,
    body_size: int = 11,
) -> dict:
    s = {}

    s["body"] = ParagraphStyle(
        "body",
        fontName=body_font, fontSize=body_size, leading=body_size * 1.75,
        textColor=INK, alignment=TA_JUSTIFY,
        spaceBefore=0, spaceAfter=10,
        firstLineIndent=18,
    )
    s["body_first"] = ParagraphStyle(
        "body_first",
        parent=s["body"],
        firstLineIndent=0,
        spaceAfter=10,
    )
    s["chapter_title"] = ParagraphStyle(
        "chapter_title",
        fontName=bold_font, fontSize=chapter_title_size, leading=int(chapter_title_size * 1.3),
        textColor=INK, alignment=TA_LEFT,
        spaceBefore=4, spaceAfter=6,
    )
    s["chapter_num"] = ParagraphStyle(
        "chapter_num",
        fontName=body_font, fontSize=10, leading=14,
        textColor=ACCENT, alignment=TA_LEFT,
        spaceBefore=0, spaceAfter=4,
        letterSpacing=2,
    )
    s["section_head"] = ParagraphStyle(
        "section_head",
        fontName=bold_font, fontSize=subsection_title_size, leading=max(subsection_title_size + 4, 20),
        textColor=INK, alignment=TA_LEFT,
        spaceBefore=18, spaceAfter=6,
    )
    s["subsection_head"] = ParagraphStyle(
        "subsection_head",
        fontName=bold_font, fontSize=max(subsection_title_size - 1, 13), leading=max(subsection_title_size + 2, 18),
        textColor=INK_LIGHT, alignment=TA_LEFT,
        spaceBefore=12, spaceAfter=4,
    )
    s["toc_title"] = ParagraphStyle(
        "toc_title",
        fontName=bold_font, fontSize=16, leading=22,
        textColor=INK, alignment=TA_LEFT,
        spaceBefore=0, spaceAfter=20,
    )
    s["toc_chapter"] = ParagraphStyle(
        "toc_chapter",
        fontName=body_font, fontSize=body_size, leading=body_size * 1.8,
        textColor=INK, alignment=TA_LEFT,
        spaceBefore=0, spaceAfter=2,
        leftIndent=0,
    )
    s["toc_num"] = ParagraphStyle(
        "toc_num",
        fontName=bold_font, fontSize=body_size, leading=body_size * 1.8,
        textColor=ACCENT, alignment=TA_LEFT,
        spaceBefore=0, spaceAfter=2,
    )
    s["meta_label"] = ParagraphStyle(
        "meta_label",
        fontName=bold_font, fontSize=8, leading=12,
        textColor=ACCENT, alignment=TA_LEFT,
        spaceBefore=0, spaceAfter=2,
        letterSpacing=1.5,
    )
    s["caption"] = ParagraphStyle(
        "caption",
        fontName=italic_font, fontSize=9, leading=13,
        textColor=CAPTION_COL, alignment=TA_LEFT,
        spaceBefore=0, spaceAfter=8,
    )
    s["book_title_size"] = book_title_size
    return s


# ── Parse helpers ──────────────────────────────────────────────────────────────

def _classify_line(line: str) -> tuple[str, str, str]:
    """
    Classify a line for PDF rendering.
    Returns (kind, num_label, title).
    kind: 'chapter' | 'section' | 'subsection' | 'separator' | 'none'
    """
    stripped = line.strip()

    # Explicit chapter markers
    m = re.match(r"^(?:Rozdział|ROZDZIAŁ|Chapter|CHAPTER|Kapitel|Capítulo)\s+(\d+)[:\.]?\s*(.+)?$",
                 stripped, re.I)
    if m:
        num = m.group(1)
        title = (m.group(2) or "").strip()
        return "chapter", f"ROZDZIAŁ {num}", title

    # Numbered list item as chapter (e.g., "1. Title")
    m = re.match(r"^(\d+)\.\s+(.{5,60})$", stripped)
    if m:
        return "chapter", f"ROZDZIAŁ {m.group(1)}", m.group(2).strip()

    # ### subsection
    m = re.match(r"^###\s+(.+)$", stripped)
    if m:
        return "subsection", "", m.group(1).strip()

    # ## section heading (NOT chapter)
    m = re.match(r"^##\s+(.+)$", stripped)
    if m:
        return "section", "", m.group(1).strip()

    # # top-level chapter heading
    m = re.match(r"^#\s+(.+)$", stripped)
    if m:
        return "chapter", "", m.group(1).strip()

    # Pipeline separator ======Title======
    m = re.match(r"^={20,}\s*$", stripped)
    if m:
        return "separator", "", ""

    m = re.match(r"^={3,}\s*(.+?)\s*=*$", stripped)
    if m:
        title = m.group(1).strip()
        if title:
            return "chapter", "", title

    return "none", "", ""


def _clean_llm_text(text: str) -> str:
    """Strip markdown / URL artifacts that LLMs add but break plain-text PDF rendering."""
    # Markdown links [text](url) → text
    text = re.sub(r'\[([^\]]*)\]\([^)]*\)', r'\1', text)
    # Standalone URLs
    text = re.sub(r'https?://\S+', '', text)
    # Bold **text** → text
    text = re.sub(r'\*\*([^*\n]+)\*\*', r'\1', text)
    # Italic *text* → text
    text = re.sub(r'\*([^*\n]+)\*', r'\1', text)
    # Inline hashtags like #keyword (not heading lines)
    text = re.sub(r'(?<!\n)(?<!\A)#(\w)', r'\1', text)
    # Horizontal rules --- / *** / ___
    text = re.sub(r'^[-*_]{3,}\s*$', '', text, flags=re.MULTILINE)
    # Collapse multiple blank lines to max two
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def _parse_manuscript(text: str) -> list[dict]:
    """
    Parse manuscript into chapters.
    Each chapter: {num, title, content}
    content: list of items — either str (paragraph) or dict {type, text}
    """
    if not text:
        return []

    text = _clean_llm_text(text)

    chapters: list[dict] = []
    current_num = ""
    current_title = ""
    current_content: list = []

    def _flush():
        if current_content or current_title:
            chapters.append({
                "num": current_num,
                "title": current_title,
                "content": [c for c in current_content if
                            (isinstance(c, str) and c.strip()) or
                            (isinstance(c, dict))],
            })

    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        kind, num, title = _classify_line(stripped)

        if kind == "separator":
            # Check for ======\nTitle\n====== pattern (pipeline chapter separator)
            j = i + 1
            while j < len(lines) and not lines[j].strip():
                j += 1
            if j < len(lines):
                candidate_title = lines[j].strip()
                k = j + 1
                while k < len(lines) and not lines[k].strip():
                    k += 1
                if (k < len(lines) and
                        re.match(r"^={3,}\s*$", lines[k].strip()) and
                        candidate_title and
                        len(candidate_title) < 120):
                    # This is a ===\nTitle\n=== chapter block
                    _flush()
                    current_num = ""
                    current_title = candidate_title
                    current_content = []
                    i = k + 1
                    continue
            i += 1
            continue
        elif kind == "chapter" and title:
            _flush()
            current_num = num
            current_title = title
            current_content = []
        elif kind in ("section", "subsection") and title:
            current_content.append({"type": kind, "text": title})
        elif stripped == "":
            # Blank line = paragraph break marker
            if current_content and current_content[-1] != "":
                current_content.append("")
        else:
            current_content.append(stripped)

        i += 1

    _flush()
    return chapters


def _locale_key(language: str | None) -> str:
    raw = (language or "").strip().lower()
    if raw.startswith("de"):
        return "de"
    if raw.startswith("en"):
        return "en"
    return "pl"


def _toc_heading(language: str | None) -> str:
    return {
        "de": "Inhaltsverzeichnis",
        "en": "Table of contents",
        "pl": "Spis treści",
    }[_locale_key(language)]


def _chapter_label(language: str | None, index: int) -> str:
    prefix = {
        "de": "Kapitel",
        "en": "Chapter",
        "pl": "Rozdział",
    }[_locale_key(language)]
    return f"{prefix} {index}"


# ── Main exporter ──────────────────────────────────────────────────────────────

class ExportService:

    # ── DOCX ──────────────────────────────────────────────────────────────────

    def build_docx(self, project: BookProject) -> bytes:
        doc = Document()

        style = doc.styles["Normal"]
        style.font.name = "Georgia"
        style.font.size = Pt(11)

        # Title page
        t = doc.add_paragraph()
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = t.add_run(project.title)
        run.bold = True
        run.font.size = Pt(28)
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

        doc.add_paragraph()
        if project.author_bio:
            author_line = project.author_bio[:80]
            meta = doc.add_paragraph(author_line)
            meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if meta.runs:
                meta.runs[0].font.size = Pt(12)

        doc.add_page_break()

        content = project.edited_text or project.manuscript_text or ""
        if content.strip():
            sections_map = [("Manuskrypt", content)]
        else:
            sections_map = [
                ("Konspekt", project.outline_text),
                ("Prompty rozdziałów", project.chapter_prompts),
                ("Draft", project.manuscript_text),
                ("Zredagowany manuskrypt", project.edited_text),
            ]

        for sec_title, sec_content in sections_map:
            if not (sec_content or "").strip():
                continue
            h = doc.add_heading(sec_title, level=1)
            h.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

            for para_text in (sec_content or "").split("\n\n"):
                para_text = para_text.strip()
                if not para_text:
                    continue
                p = doc.add_paragraph(para_text)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                if p.runs:
                    p.runs[0].font.size = Pt(11)
            doc.add_page_break()

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ── PDF ───────────────────────────────────────────────────────────────────

    def build_pdf(self, project: BookProject) -> bytes:
        _register_fonts()
        self._current_language = getattr(project, "language", "pl")

        font_family = _coerce_font_family(getattr(project, "pdf_font_family", "Georgia"))
        trim_size = _coerce_trim_size(getattr(project, "pdf_trim_size", "6x9"))
        book_title_size = _coerce_font_size(
            getattr(project, "pdf_book_title_size", 30),
            default=30,
            minimum=18,
            maximum=60,
        )
        chapter_title_size = _coerce_font_size(
            getattr(project, "pdf_chapter_title_size", getattr(project, "pdf_heading_size", 23)),
            default=23,
            minimum=14,
            maximum=48,
        )
        subsection_title_size = _coerce_font_size(
            getattr(project, "pdf_subchapter_title_size", 17),
            default=17,
            minimum=12,
            maximum=36,
        )
        body_size = _coerce_font_size(
            getattr(project, "pdf_body_size", 11),
            default=11,
            minimum=8,
            maximum=18,
        )

        body_font, bold_font, italic_font = _set_active_font(font_family)
        styles = _build_styles(
            body_font,
            bold_font,
            italic_font,
            book_title_size=book_title_size,
            chapter_title_size=chapter_title_size,
            subsection_title_size=subsection_title_size,
            body_size=body_size,
        )

        buf = BytesIO()
        page_w, page_h = _trim_page_size(trim_size)
        margin_inner = 0.72 * inch if trim_size == "6x9" else 25 * mm
        margin_outer = 0.62 * inch if trim_size == "6x9" else 20 * mm

        doc = SimpleDocTemplate(
            buf,
            pagesize=(page_w, page_h),
            leftMargin=margin_inner,
            rightMargin=margin_outer,
            topMargin=(0.82 * inch if trim_size == "6x9" else 22 * mm + 8 * mm),
            bottomMargin=(0.72 * inch if trim_size == "6x9" else 22 * mm + 6 * mm),
            title=self._pdf_title(project),
            author=self._extract_author_name(project),
        )

        story: list = []

        # ── 1. Title page ──────────────────────────────────────────────
        story.extend(self._build_title_page(project, styles, page_w, page_h))
        story.append(PageBreak())

        # ── 2. Table of Contents ───────────────────────────────────────
        manuscript = (project.edited_text or project.manuscript_text or "").strip()
        chapters = _parse_manuscript(manuscript) if manuscript else []
        include_toc = _coerce_bool(getattr(project, "pdf_include_toc", True), default=True)
        show_page_numbers = _coerce_bool(getattr(project, "pdf_show_page_numbers", True), default=True)

        if chapters and include_toc:
            story.extend(self._build_toc(chapters, styles, page_w))
            story.append(PageBreak())

        if chapters:
            # ── 3. Chapters ────────────────────────────────────────────
            story.extend(self._build_chapters(chapters, styles, page_w, margin_inner, margin_outer))
        elif manuscript:
            story.append(Paragraph("Manuskrypt", styles["section_head"]))
            story.append(OrnamentalRule(page_w - margin_inner - margin_outer))
            story.append(Spacer(1, 6 * mm))
            for block in manuscript.split("\n\n"):
                block = block.strip()
                if block:
                    story.append(Paragraph(self._escape(block), styles["body"]))

        def _on_page(canvas, doc):
            canvas.saveState()
            _page_header_footer(
                canvas,
                doc,
                self._pdf_title(project),
                self._extract_author_name(project),
                header_font=italic_font,
                footer_font=body_font,
                show_page_numbers=show_page_numbers,
                accent_color=ACCENT,
            )
            canvas.restoreState()

        doc.build(story, onFirstPage=_cover_page_bg, onLaterPages=_on_page)
        return buf.getvalue()

    # ── Title page builder ────────────────────────────────────────────────────

    def _build_title_page(self, project: BookProject, styles: dict, pw: float, ph: float) -> list:
        elements: list = []

        elements.append(Spacer(1, ph * 0.22))

        title_style = ParagraphStyle(
            "cover_title",
            fontName=styles["chapter_title"].fontName,
            fontSize=styles["book_title_size"],
            leading=int(styles["book_title_size"] * 1.22),
            textColor=INK,
            alignment=TA_LEFT,
        )
        elements.append(Paragraph(self._escape(self._pdf_title(project)), title_style))
        subtitle = (getattr(project, "pdf_subtitle", "") or "").strip()
        if subtitle:
            subtitle_style = ParagraphStyle(
                "cover_subtitle",
                fontName=styles["body"].fontName,
                fontSize=max(styles["body"].fontSize + 2, 13),
                leading=max(styles["body"].leading, 18),
                textColor=INK_LIGHT,
                alignment=TA_LEFT,
            )
            elements.append(Spacer(1, 3 * mm))
            elements.append(Paragraph(self._escape(subtitle), subtitle_style))
        elements.append(Spacer(1, 6 * mm))
        elements.append(HRFlowable(width="100%", thickness=2, color=ACCENT, spaceAfter=6 * mm))

        author_name = self._extract_author_name(project)
        if author_name:
            author_style = ParagraphStyle(
                "cover_author",
                fontName=styles["body"].fontName,
                fontSize=14,
                leading=20,
                textColor=colors.HexColor("#4b5563"),
                alignment=TA_LEFT,
            )
            elements.append(Paragraph(self._escape(author_name), author_style))

        return elements

    # ── TOC builder ───────────────────────────────────────────────────────────

    def _build_toc(self, chapters: list[dict], styles: dict, pw: float) -> list:
        elements: list = []

        elements.append(Spacer(1, 6 * mm))
        elements.append(Paragraph(_toc_heading(getattr(self, "_current_language", "pl")), styles["toc_title"]))
        elements.append(HRFlowable(width="40%", thickness=1, color=ACCENT_SOFT, spaceAfter=8 * mm))

        for idx, ch in enumerate(chapters, 1):
            num_label = ch["num"] or _chapter_label(getattr(self, "_current_language", "pl"), idx)
            title = ch["title"] or ""
            if title:
                row_text = f"<b>{self._escape(num_label)}</b>  {self._escape(title)}"
            else:
                row_text = f"<b>{self._escape(num_label)}</b>"
            elements.append(Paragraph(row_text, styles["toc_chapter"]))

        return elements

    # ── Chapter builder ────────────────────────────────────────────────────────

    def _build_chapters(self, chapters: list[dict], styles: dict,
                        pw: float, ml: float, mr: float) -> list:
        elements: list = []
        for idx, ch in enumerate(chapters):
            if idx > 0:
                elements.append(PageBreak())

            # Chapter number label
            if ch["num"]:
                elements.append(Paragraph(ch["num"], styles["chapter_num"]))
                elements.append(Spacer(1, 2 * mm))

            # Chapter title
            if ch["title"]:
                elements.append(Paragraph(self._escape(ch["title"]), styles["chapter_title"]))

            # Spacer after title (the "przeskok do następnej linii" the client requested)
            elements.append(Spacer(1, 4 * mm))
            elements.append(OrnamentalRule(pw - ml - mr))
            elements.append(Spacer(1, 6 * mm))

            # Content items
            first = True
            for item in ch["content"]:
                if isinstance(item, dict):
                    if item["type"] == "section":
                        elements.append(Paragraph(self._escape(item["text"]), styles["section_head"]))
                        # Line break after section heading before continuing text
                        elements.append(Spacer(1, 3 * mm))
                        first = True
                    elif item["type"] == "subsection":
                        elements.append(Paragraph(self._escape(item["text"]), styles["subsection_head"]))
                        elements.append(Spacer(1, 2 * mm))
                        first = True
                elif isinstance(item, str):
                    if not item.strip():
                        elements.append(Spacer(1, 4 * mm))
                        first = True
                    else:
                        s = styles["body_first"] if first else styles["body"]
                        elements.append(Paragraph(self._escape(item), s))
                        first = False

        return elements

    # ── Helpers ────────────────────────────────────────────────────────────────

    @staticmethod
    def _escape(text: str) -> str:
        return (
            text
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )

    @staticmethod
    def _extract_author_name(project: BookProject) -> str:
        custom = (getattr(project, "pdf_author_name", "") or "").strip()
        if custom:
            return custom[:80]
        bio = (project.author_bio or "").strip()
        if not bio:
            return ""
        # Take first sentence or up to first comma/dash/newline
        name = re.split(r'[,\-–—\n]', bio)[0].strip()
        return name[:80]

    @staticmethod
    def _pdf_title(project: BookProject) -> str:
        return ((getattr(project, "pdf_title_override", "") or "").strip() or project.title).strip()

    def _sections(self, project: BookProject):
        all_sections = [
            ("Konspekt", project.outline_text),
            ("Prompty rozdziałów", project.chapter_prompts),
            ("Draft książki", project.manuscript_text),
            ("Zredagowany manuskrypt", project.edited_text),
            ("Opis Amazon SEO", project.seo_description),
            ("Brief okładki", project.cover_brief),
            ("Checklista publikacji", project.publish_checklist),
        ]
        return [(t, c) for t, c in all_sections if (c or "").strip()]


# ── Page decorators ────────────────────────────────────────────────────────────

def _cover_page_bg(canvas, doc):
    canvas.saveState()
    w, h = doc.pagesize
    canvas.setFillColor(CREAM)
    canvas.rect(0, h - 50 * mm, w, 50 * mm, fill=1, stroke=0)
    canvas.setFillColor(ACCENT)
    canvas.rect(0, 0, 4, h, fill=1, stroke=0)
    canvas.setFillColor(CREAM)
    canvas.rect(w - 30 * mm, 0, 30 * mm, 30 * mm, fill=1, stroke=0)
    canvas.restoreState()


def _page_header_footer(
    canvas,
    doc,
    title: str,
    author: str = "",
    *,
    header_font: str | None = None,
    footer_font: str | None = None,
    show_page_numbers: bool = True,
    accent_color=ACCENT,
):
    if doc.page <= 1:
        return
    w, h = doc.pagesize
    ml, mr = doc.leftMargin, doc.rightMargin

    header_font = header_font or _set_active_font("auto")[2] or "Helvetica-Oblique"
    footer_font = footer_font or _set_active_font("auto")[0] or "Helvetica"

    canvas.setStrokeColor(colors.HexColor("#e2e0db"))
    canvas.setLineWidth(0.4)

    # Header
    canvas.line(ml, h - 18 * mm, w - mr, h - 18 * mm)
    canvas.setFont(header_font, 8)
    canvas.setFillColor(colors.HexColor("#9ca3af"))
    if doc.page % 2 == 0:
        canvas.drawString(ml, h - 15 * mm, _canvas_safe_text(title, 60))
    else:
        right_label = _canvas_safe_text(author, 40) if author else "Book Factory"
        canvas.drawRightString(w - mr, h - 15 * mm, right_label)

    # Accent left bar
    canvas.setFillColor(accent_color)
    canvas.setLineWidth(2)
    canvas.line(0, 0, 0, h)

    # Footer
    canvas.setLineWidth(0.4)
    canvas.setStrokeColor(colors.HexColor("#e2e0db"))
    canvas.line(ml, 16 * mm, w - mr, 16 * mm)
    if show_page_numbers:
        canvas.setFont(footer_font, 9)
        canvas.setFillColor(colors.HexColor("#9ca3af"))
        canvas.drawCentredString(w / 2, 12 * mm, str(doc.page - 1))


export_service = ExportService()
